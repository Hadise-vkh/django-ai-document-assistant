from django.db import models
from docx import Document as DocxDocument


class Document(models.Model):
    title = models.CharField(max_length=255)
    file = models.FileField(upload_to="documents/")
    content = models.TextField(blank=True)
    created_at = models.DateTimeField(auto_now_add=True)

    def save(self, *args, **kwargs):
        super().save(*args, **kwargs)

        if self.file:
            doc = DocxDocument(self.file.path)

            text = "\n".join(
                paragraph.text
                for paragraph in doc.paragraphs
            )

            self.content = text

            super().save(update_fields=["content"])

    def __str__(self):
        return self.title